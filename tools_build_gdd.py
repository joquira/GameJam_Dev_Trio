from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/GDD_Atras_da_Porta.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_reference_box(doc, title, color, bullets):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.35, 5.0])
    left, right = table.rows[0].cells
    set_cell_shading(left, color)
    left.text = "Print /\nreferencia"
    for paragraph in left.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    right.text = title
    right.paragraphs[0].runs[0].bold = True
    for bullet in bullets:
        p = right.add_paragraph(bullet, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in (
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GDD - Atras da Porta")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph("AC1 Game Jam | Godot 4 | Plataforma 2D | Tema: What's behind the door")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. Visao geral", 1)
    add_heading(doc, "Nome do jogo", 2)
    doc.add_paragraph("Atras da Porta")
    add_heading(doc, "Premissa", 2)
    doc.add_paragraph(
        "O jogador controla Lio, um aventureiro em pixel art retirado do spritesheet characters.png, preso em uma ruina simples de plataforma 2D. "
        "Para vencer, ele precisa atravessar a fase, coletar dois baus flutuantes e encostar na porta final."
    )

    add_heading(doc, "Referencias de gameplay", 2)
    add_reference_box(
        doc,
        "Celeste",
        "2E74B5",
        [
            "Inspiracao para controle responsivo, fases curtas e desafio crescente.",
            "Absorvido: pulo preciso, plataformas com zonas de respiro antes de trechos mais dificeis.",
        ],
    )
    add_reference_box(
        doc,
        "Hollow Knight",
        "1F4D78",
        [
            "Inspiracao para atmosfera misteriosa e descoberta progressiva do mundo.",
            "Absorvido: a porta final como promessa visual de algo escondido.",
        ],
    )
    add_reference_box(
        doc,
        "Super Mario World",
        "2B8A3E",
        [
            "Inspiracao para leitura clara da fase e plataformas com dificuldade gradual.",
            "Absorvido: inicio seguro, obstaculos apresentados um de cada vez e final bem marcado.",
        ],
    )

    add_heading(doc, "Divisao de tarefas", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [2.0, 2.0, 2.3])
    headers = ["Area", "Responsavel", "Entrega"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        cell.paragraphs[0].runs[0].bold = True
    for area, resp, entrega in (
        ("Documentacao", "Integrante 1", "GDD, README e organizacao do repositorio"),
        ("Personagem", "Integrante 2", "Player com sprites de characters.png, animacoes idle/walk/run/jump e controle"),
        ("Level", "Integrante 3", "TileMap com blocos do sheet.png, inimigos, baus coletaveis, porta final e curva de dificuldade"),
    ):
        row = table.add_row().cells
        row[0].text = area
        row[1].text = resp
        row[2].text = entrega

    add_heading(doc, "2. Personagem principal", 1)
    add_heading(doc, "Nome e descricao", 2)
    doc.add_paragraph(
        "Lio e um pequeno aventureiro de roupa laranja em pixel art. Ele acorda dentro de uma ruina sem lembrar como entrou. "
        "A historia dele gira em torno de atravessar a fase e descobrir o que existe atras da porta final."
    )
    add_heading(doc, "Movimentacao", 2)
    add_bullets(
        doc,
        [
            "Velocidade andando: controle medio para permitir saltos precisos.",
            "Corrida com Shift: mais rapida, usada para vencer lacunas maiores.",
            "Pulo alto e responsivo: o personagem deve parecer agil, porque o foco e plataforma 2D.",
            "Visual atual: sprite 2D do arquivo assets/sprites/characters.png, ampliado em escala pixel art sem filtro borrado.",
        ],
    )
    add_heading(doc, "Habilidades especiais", 2)
    doc.add_paragraph(
        "A versao atual nao usa habilidade especial. O jogo foi simplificado para movimento basico, corrida, pulo, coleta de baus e chegada na porta final."
    )

    add_heading(doc, "3. Mecanicas", 1)
    add_heading(doc, "Mecanica central", 2)
    doc.add_paragraph(
        "Explorar uma fase de plataforma, saltar entre blocos do tileset, evitar inimigos, coletar baus flutuantes e chegar ate a porta final."
    )
    add_heading(doc, "Objetivo", 2)
    doc.add_paragraph("Chegar ao fim da fase depois de coletar 2 baus e encostar na porta final.")
    add_heading(doc, "Obstaculos e desafios", 2)
    add_bullets(
        doc,
        [
            "Buracos entre plataformas.",
            "Inimigos patrulhando trechos de acao.",
            "Saltos que exigem corrida e precisao.",
            "Porta final so vence o jogo se o jogador ja tiver coletado os 2 baus.",
        ],
    )
    add_heading(doc, "Coletaveis", 2)
    doc.add_paragraph(
        "Existem 2 baus flutuantes, usando os sprites de bau do arquivo assets/tiles/sheet.png. Eles funcionam como chave narrativa e mecanica para liberar a porta final."
    )
    add_heading(doc, "Conexao com o tema", 2)
    doc.add_paragraph(
        "O tema aparece de forma simples: a porta final representa o misterio do que existe atras dela. "
        "O jogador precisa explorar a fase e coletar os baus antes de descobrir o que ha atras da porta."
    )

    add_heading(doc, "4. Inimigos", 1)
    enemy_table = doc.add_table(rows=1, cols=3)
    enemy_table.style = "Table Grid"
    set_table_width(enemy_table, [1.5, 2.4, 2.4])
    for idx, header in enumerate(["Inimigo", "Descricao visual", "Comportamento"]):
        cell = enemy_table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        cell.paragraphs[0].runs[0].bold = True
    row = enemy_table.add_row().cells
    row[0].text = "Sentinela Vermelha"
    row[1].text = "Inimigo em pixel art retirado do spritesheet characters.png, usando escala menor que o personagem principal."
    row[2].text = "Patrulha horizontalmente entre dois limites. Se encostar no jogador, ele volta ao inicio."

    add_heading(doc, "5. Level design da fase", 1)
    add_heading(doc, "Descricao geral", 2)
    doc.add_paragraph(
        "A fase se passa em uma ruina escura, com uma porta final, baus flutuantes e plataformas feitas apenas com os blocos de chao do sheet.png. "
        "A atmosfera e de misterio e descoberta, com visual 2D em pixel art."
    )
    add_heading(doc, "Estrutura", 2)
    doc.add_paragraph(
        "A fase e linear com pequenos desvios verticais para coletaveis. Ela possui zonas de respiro, zonas de acao e um final claro."
    )
    add_heading(doc, "Curva de dificuldade", 2)
    add_bullets(
        doc,
        [
            "Zona 1 - respiro: chao largo e poucos saltos para ensinar movimento, corrida e pulo.",
            "Zona 2 - acao: primeiras lacunas, um inimigo e o primeiro bau flutuante.",
            "Zona 3 - aumento: plataformas mais altas, segundo inimigo e segundo bau flutuante.",
            "Final - respiro curto: porta final e saida visivel para concluir a fase.",
        ],
    )
    add_heading(doc, "Condicao de vitoria", 2)
    doc.add_paragraph("Coletar os 2 baus e encostar na porta final.")

    add_heading(doc, "6. Entregaveis funcionais implementados", 1)
    add_bullets(
        doc,
        [
            "Projeto Godot 4 com cena principal em scenes/main.tscn.",
            "Personagem principal com sprites 2D e animacoes funcionais: idle, walk, run e jump.",
            "Mapa jogavel com TileMap visual feito com os blocos de chao do sheet.png, plataformas, inicio e fim definidos.",
            "Fase navegavel do inicio ao fim.",
            "Porta final conectada ao tema What's behind the door.",
            "Baus flutuantes como coletaveis, inimigos em sprite 2D e HUD simples.",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.text = "Atras da Porta - AC1 Game Jam"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)


if __name__ == "__main__":
    build()
